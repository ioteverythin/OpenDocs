"""Generate a Word (.docx) technical report from a DocumentModel."""

from __future__ import annotations

from pathlib import Path

from docx import Document as DocxDocument
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Emu, Inches, Mm, Pt, RGBColor

from ..core.models import (
    BlockquoteBlock,
    CodeBlock,
    ContentBlock,
    DocumentModel,
    GenerationResult,
    HeadingBlock,
    ImageBlock,
    InlineSpan,
    ListBlock,
    MermaidBlock,
    OutputFormat,
    ParagraphBlock,
    Section,
    TableBlock,
    ThematicBreakBlock,
)
from .base import BaseGenerator
from .styles import Colors, Fonts, Layout


def _hex(rgb: tuple[int, int, int]) -> str:
    return "{:02X}{:02X}{:02X}".format(*rgb)


def _set_cell_shading(cell, color: tuple[int, int, int]) -> None:
    """Apply background shading to a table cell."""
    tc_pr = cell._element.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), _hex(color))
    shading.set(qn("w:val"), "clear")
    tc_pr.append(shading)


def _set_cell_margins(cell, top=40, bottom=40, left=80, right=80):
    """Set cell margins in twips."""
    tc_pr = cell._element.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for side, val in [("top", top), ("bottom", bottom), ("start", left), ("end", right)]:
        el = OxmlElement(f"w:{side}")
        el.set(qn("w:w"), str(val))
        el.set(qn("w:type"), "dxa")
        tc_mar.append(el)
    tc_pr.append(tc_mar)


def _add_bottom_border(paragraph, color: tuple[int, int, int], size: int = 6):
    """Add a colored bottom border to a paragraph."""
    p_pr = paragraph._element.get_or_add_pPr()
    p_bdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "4")
    bottom.set(qn("w:color"), _hex(color))
    p_bdr.append(bottom)
    p_pr.append(p_bdr)


def _set_paragraph_shading(paragraph, color: tuple[int, int, int]):
    """Apply background shading to an entire paragraph."""
    p_pr = paragraph._element.get_or_add_pPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), _hex(color))
    shading.set(qn("w:val"), "clear")
    p_pr.append(shading)


def _add_hyperlink(paragraph, text: str, url: str):
    """Add a clickable hyperlink run to a paragraph."""
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    # Blue underlined style
    color_el = OxmlElement("w:color")
    color_el.set(qn("w:val"), "1565C0")
    rPr.append(color_el)
    u_el = OxmlElement("w:u")
    u_el.set(qn("w:val"), "single")
    rPr.append(u_el)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), str(Fonts.BODY_SIZE_PT * 2))  # half-points
    rPr.append(sz)
    run_el.append(rPr)

    t_el = OxmlElement("w:t")
    t_el.set(qn("xml:space"), "preserve")
    t_el.text = text
    run_el.append(t_el)

    hyperlink.append(run_el)
    paragraph._element.append(hyperlink)


def _add_rich_runs(paragraph, spans: list[InlineSpan]):
    """Render a list of InlineSpans into a paragraph with hyperlinks and formatting."""
    for span in spans:
        if not span.text:
            continue
        if span.is_link:
            _add_hyperlink(paragraph, span.text, span.url)
        else:
            run = paragraph.add_run(span.text)
            run.font.size = Pt(Fonts.BODY_SIZE_PT)
            run.font.name = Fonts.BODY
            if span.bold:
                run.bold = True
            if span.italic:
                run.italic = True
            if span.code:
                run.font.name = Fonts.CODE
                run.font.size = Pt(Fonts.CODE_SIZE_PT)


class WordGenerator(BaseGenerator):
    """Generates a beautifully styled ``.docx`` Word document."""

    format = OutputFormat.WORD

    def generate(self, doc: DocumentModel, output_dir: Path) -> GenerationResult:
        output_dir = self._ensure_dir(output_dir)
        fname = self._safe_filename(doc.metadata.repo_name or "report", "docx")
        output_path = output_dir / fname
        # Track which mermaid diagram index we're on during rendering
        self._mermaid_index = 0

        try:
            docx = self._build(doc)
            docx.save(str(output_path))
            return GenerationResult(format=self.format, output_path=output_path)
        except Exception as exc:
            return GenerationResult(
                format=self.format, output_path=output_path, success=False, error=str(exc)
            )

    # ------------------------------------------------------------------
    # Builder
    # ------------------------------------------------------------------

    def _build(self, doc: DocumentModel) -> DocxDocument:
        docx = DocxDocument()

        # -- Page setup --
        for section in docx.sections:
            section.top_margin = Inches(Layout.PAGE_MARGIN_INCHES * 0.8)
            section.bottom_margin = Inches(Layout.PAGE_MARGIN_INCHES * 0.8)
            section.left_margin = Inches(Layout.PAGE_MARGIN_INCHES)
            section.right_margin = Inches(Layout.PAGE_MARGIN_INCHES)

        # -- Default body style --
        style = docx.styles["Normal"]
        font = style.font
        font.name = Fonts.BODY
        font.size = Pt(Fonts.BODY_SIZE_PT)
        font.color.rgb = RGBColor(*Colors.TEXT)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.2

        # -- Customize heading styles --
        self._setup_heading_styles(docx)

        # -- Add headers/footers from template variables --
        if self.tvars.has_values:
            self._add_header_footer(docx)

        # -- Title page --
        self._add_title_page(docx, doc)

        # -- Table of Contents --
        self._add_toc_page(docx)

        # -- Body sections --
        for section in doc.sections:
            self._render_section(docx, section)

        # -- Knowledge Graph page (if available) --
        if self.kg and self.kg.entities:
            self._add_knowledge_graph_page(docx)

        # -- Metadata footer page --
        self._add_metadata_page(docx, doc)

        return docx

    # ------------------------------------------------------------------
    # Style setup
    # ------------------------------------------------------------------

    def _setup_heading_styles(self, docx: DocxDocument) -> None:
        """Customize built-in heading styles for a polished look."""
        heading_config = [
            ("Heading 1", Fonts.H1_SIZE_PT, Colors.PRIMARY_DARK, 24, 10, True),
            ("Heading 2", Fonts.H2_SIZE_PT, Colors.HEADING, 18, 8, True),
            ("Heading 3", Fonts.H3_SIZE_PT, Colors.HEADING, 14, 6, False),
            ("Heading 4", Fonts.H4_SIZE_PT, Colors.SECONDARY, 10, 4, False),
        ]
        for name, size, color, space_before, space_after, bold in heading_config:
            if name in [s.name for s in docx.styles]:
                style = docx.styles[name]
                style.font.name = Fonts.HEADING
                style.font.size = Pt(size)
                style.font.color.rgb = RGBColor(*color)
                style.font.bold = bold
                style.paragraph_format.space_before = Pt(space_before)
                style.paragraph_format.space_after = Pt(space_after)
                style.paragraph_format.keep_with_next = True

    # ------------------------------------------------------------------
    # Headers & Footers (from template variables)
    # ------------------------------------------------------------------

    def _add_header_footer(self, docx: DocxDocument) -> None:
        """Add document headers and footers from template variables."""
        for section in docx.sections:
            # --- Header ---
            header_text = self.tvars.header_text
            if header_text:
                header = section.header
                header.is_linked_to_previous = False
                p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                run = p.add_run(header_text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(*Colors.MUTED)
                run.font.name = Fonts.BODY
                _add_bottom_border(p, Colors.ACCENT, size=4)

            # --- Footer ---
            footer_text = self.tvars.footer_text
            if footer_text:
                footer = section.footer
                footer.is_linked_to_previous = False
                p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = p.add_run(footer_text)
                run.font.size = Pt(8)
                run.font.color.rgb = RGBColor(*Colors.MUTED)
                run.font.name = Fonts.BODY

    # ------------------------------------------------------------------
    # Title page
    # ------------------------------------------------------------------

    def _add_title_page(self, docx: DocxDocument, doc: DocumentModel) -> None:
        # Top accent bar
        bar = docx.add_paragraph()
        bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = bar.add_run("━" * 50)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*Colors.ACCENT)

        # Spacer
        for _ in range(3):
            spacer = docx.add_paragraph()
            spacer.paragraph_format.space_after = Pt(0)

        # Decorative icon text
        icon_p = docx.add_paragraph()
        icon_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = icon_p.add_run("")
        run.font.size = Pt(48)

        # Title
        title = docx.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title.paragraph_format.space_before = Pt(12)
        title.paragraph_format.space_after = Pt(4)
        # Use template project_name if available, else repo_name
        title_text = self.tvars.project_name or doc.metadata.repo_name or "Technical Report"
        run = title.add_run(title_text)
        run.font.size = Pt(Fonts.TITLE_SIZE_PT)
        run.font.color.rgb = RGBColor(*Colors.PRIMARY_DARK)
        run.bold = True
        run.font.name = Fonts.HEADING

        # Version badge (from template vars)
        if self.tvars.version:
            ver_p = docx.add_paragraph()
            ver_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = ver_p.add_run(f"v{self.tvars.version}")
            run.font.size = Pt(11)
            run.font.color.rgb = RGBColor(*Colors.ACCENT)
            run.bold = True
            run.font.name = Fonts.BODY

        # Accent underline
        line = docx.add_paragraph()
        line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = line.add_run("▬" * 20)
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(*Colors.ACCENT)

        # Subtitle / description
        if doc.metadata.description:
            sub = docx.add_paragraph()
            sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sub.paragraph_format.space_before = Pt(12)
            run = sub.add_run(doc.metadata.description[:250])
            run.font.size = Pt(13)
            run.font.color.rgb = RGBColor(*Colors.MUTED)
            run.italic = True
            run.font.name = Fonts.BODY

        # Spacer
        for _ in range(3):
            docx.add_paragraph()

        # Metadata box
        info_items = [
            ("Repository", doc.metadata.repo_name),
            ("Source", doc.metadata.repo_url or doc.metadata.source_path),
            ("Generated", doc.metadata.generated_at[:19] if doc.metadata.generated_at else ""),
            ("Tool", "opendocs v0.4.1"),
        ]
        # Add template variable items
        if self.tvars.author:
            info_items.insert(0, ("Author", self.tvars.author))
        if self.tvars.organisation:
            info_items.insert(0, ("Organisation", self.tvars.organisation))
        if self.tvars.version:
            info_items.append(("Version", f"v{self.tvars.version}"))
        if self.tvars.date:
            info_items.append(("Date", self.tvars.date))
        if self.tvars.confidentiality:
            info_items.append(("Classification", self.tvars.confidentiality))
        if self.tvars.department:
            info_items.append(("Department", self.tvars.department))
        info_table = docx.add_table(rows=len(info_items), cols=2)
        info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(info_items):
            left_cell = info_table.cell(i, 0)
            right_cell = info_table.cell(i, 1)
            left_cell.text = label
            right_cell.text = value or "—"
            _set_cell_shading(left_cell, Colors.BG_LIGHT)
            for p in left_cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(*Colors.HEADING)
                    r.font.name = Fonts.BODY
            for p in right_cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(*Colors.TEXT)
                    r.font.name = Fonts.BODY
            _set_cell_margins(left_cell)
            _set_cell_margins(right_cell)

        # Bottom accent bar
        docx.add_paragraph()
        bar = docx.add_paragraph()
        bar.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = bar.add_run("━" * 50)
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(*Colors.ACCENT)

        docx.add_page_break()

    # ------------------------------------------------------------------
    # Table of Contents
    # ------------------------------------------------------------------

    def _add_toc_page(self, docx: DocxDocument) -> None:
        toc_heading = docx.add_paragraph()
        toc_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = toc_heading.add_run("Table of Contents")
        run.font.size = Pt(Fonts.H1_SIZE_PT)
        run.font.color.rgb = RGBColor(*Colors.PRIMARY_DARK)
        run.bold = True
        run.font.name = Fonts.HEADING
        _add_bottom_border(toc_heading, Colors.ACCENT, size=8)

        docx.add_paragraph()
        note = docx.add_paragraph()
        run = note.add_run("To update: Right-click → Update Field  (or  References → Update Table)")
        run.font.size = Pt(Fonts.CAPTION_SIZE_PT)
        run.font.color.rgb = RGBColor(*Colors.MUTED)
        run.italic = True

        docx.add_paragraph()
        docx.add_page_break()

    # ------------------------------------------------------------------
    # Section rendering
    # ------------------------------------------------------------------

    def _render_section(self, docx: DocxDocument, section: Section) -> None:
        level = min(section.level, 4)
        heading_p = docx.add_heading(section.title, level=level)

        # Add accent bottom border to H1 and H2
        if level <= 2:
            _add_bottom_border(heading_p, Colors.ACCENT if level == 1 else Colors.PRIMARY_LIGHT)

        # -- LLM-polished prose (if available) ---------------------------
        rewritten = (
            self.kg.llm_sections.get(section.title, "")
            if self.kg and self.kg.llm_sections
            else ""
        )
        if rewritten:
            # Add LLM-rewritten narrative as a styled overview paragraph,
            # then still render the original blocks (code, tables, etc.)
            p = docx.add_paragraph()
            run = p.add_run(rewritten)
            run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(*Colors.TEXT)
            p.paragraph_format.space_after = Pt(8)
            p.paragraph_format.space_before = Pt(2)

        for block in section.blocks:
            self._render_block(docx, block)

        for sub in section.subsections:
            self._render_section(docx, sub)

    def _render_block(self, docx: DocxDocument, block: ContentBlock) -> None:
        if isinstance(block, ParagraphBlock):
            p = docx.add_paragraph()
            if block.spans:
                _add_rich_runs(p, block.spans)
            else:
                p.add_run(block.text)
            p.paragraph_format.space_after = Pt(6)

        elif isinstance(block, CodeBlock):
            self._render_code_block(docx, block)

        elif isinstance(block, MermaidBlock):
            self._render_mermaid_block(docx, block)

        elif isinstance(block, TableBlock):
            self._render_table(docx, block)

        elif isinstance(block, ListBlock):
            self._render_list(docx, block)

        elif isinstance(block, BlockquoteBlock):
            self._render_blockquote(docx, block)

        elif isinstance(block, ImageBlock):
            # Try to embed actual downloaded image
            img_path = self.image_cache.get_external(block.src) if self.image_cache else None
            if img_path and img_path.exists():
                try:
                    # Caption
                    if block.alt:
                        cap = docx.add_paragraph()
                        cap.paragraph_format.space_after = Pt(2)
                        run = cap.add_run(f"  {block.alt}")
                        run.font.size = Pt(Fonts.CAPTION_SIZE_PT)
                        run.font.color.rgb = RGBColor(*Colors.PRIMARY)
                        run.italic = True

                    docx.add_picture(str(img_path), width=Inches(5.5))
                    last_p = docx.paragraphs[-1]
                    last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    last_p.paragraph_format.space_after = Pt(8)
                except Exception:
                    # Image file may be corrupt or unsupported format
                    p = docx.add_paragraph()
                    run = p.add_run(f"[Image: {block.alt or block.src}]")
                    run.italic = True
                    run.font.color.rgb = RGBColor(*Colors.MUTED)
            else:
                p = docx.add_paragraph()
                run = p.add_run(f"[Image: {block.alt or block.src}]")
                run.italic = True
                run.font.color.rgb = RGBColor(*Colors.INFO)
                run.font.size = Pt(Fonts.CAPTION_SIZE_PT)

        elif isinstance(block, ThematicBreakBlock):
            hr = docx.add_paragraph()
            hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = hr.add_run("─" * 60)
            run.font.color.rgb = RGBColor(*Colors.TABLE_BORDER)
            run.font.size = Pt(8)
            hr.paragraph_format.space_before = Pt(8)
            hr.paragraph_format.space_after = Pt(8)

    # ------------------------------------------------------------------
    # Rich block renderers
    # ------------------------------------------------------------------

    def _render_code_block(self, docx: DocxDocument, block: CodeBlock) -> None:
        """Render a code block with dark background and language label."""
        # Language label
        if block.language:
            label_p = docx.add_paragraph()
            label_p.paragraph_format.space_before = Pt(8)
            label_p.paragraph_format.space_after = Pt(0)
            run = label_p.add_run(f"  {block.language.upper()}")
            run.font.size = Pt(Fonts.SMALL_SIZE_PT)
            run.font.color.rgb = RGBColor(*Colors.WHITE)
            run.bold = True
            run.font.name = Fonts.CODE
            _set_paragraph_shading(label_p, Colors.CODE_BG)

        # Code content in a shaded paragraph
        code_p = docx.add_paragraph()
        code_p.paragraph_format.space_before = Pt(0 if block.language else 8)
        code_p.paragraph_format.space_after = Pt(10)
        _set_paragraph_shading(code_p, Colors.CODE_BG)

        for line in block.code.rstrip("\n").split("\n"):
            run = code_p.add_run(line + "\n")
            run.font.name = Fonts.CODE
            run.font.size = Pt(Fonts.CODE_SIZE_PT)
            run.font.color.rgb = RGBColor(*Colors.CODE_TEXT)

    def _render_mermaid_block(self, docx: DocxDocument, block: MermaidBlock) -> None:
        """Render a Mermaid diagram — embed PNG image if available, else show code."""
        idx = self._mermaid_index
        self._mermaid_index += 1

        # Try to embed rendered PNG
        img_path = self.image_cache.get_mermaid(idx) if self.image_cache else None
        if img_path and img_path.exists():
            label_p = docx.add_paragraph()
            label_p.paragraph_format.space_after = Pt(2)
            run = label_p.add_run("DIAGRAM")
            run.font.size = Pt(Fonts.SMALL_SIZE_PT)
            run.font.color.rgb = RGBColor(*Colors.WHITE)
            run.bold = True
            run.font.name = Fonts.CODE
            _set_paragraph_shading(label_p, Colors.INFO)

            # Add the image, scaled to fit page width
            docx.add_picture(str(img_path), width=Inches(5.8))
            last_p = docx.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_p.paragraph_format.space_after = Pt(8)
            return

        # Fallback: show raw code
        label_p = docx.add_paragraph()
        label_p.paragraph_format.space_after = Pt(2)
        run = label_p.add_run("MERMAID DIAGRAM")
        run.font.size = Pt(Fonts.SMALL_SIZE_PT)
        run.font.color.rgb = RGBColor(*Colors.WHITE)
        run.bold = True
        run.font.name = Fonts.CODE
        _set_paragraph_shading(label_p, Colors.INFO)

        code_p = docx.add_paragraph()
        code_p.paragraph_format.space_before = Pt(0)
        code_p.paragraph_format.space_after = Pt(8)
        _set_paragraph_shading(code_p, Colors.CODE_BG)

        for line in block.code.rstrip("\n").split("\n"):
            run = code_p.add_run(line + "\n")
            run.font.name = Fonts.CODE
            run.font.size = Pt(Fonts.CODE_SIZE_PT)
            run.font.color.rgb = RGBColor(*Colors.CODE_TEXT)

    def _render_list(self, docx: DocxDocument, block: ListBlock) -> None:
        """Render a list with colored bullet markers."""
        for i, item in enumerate(block.items):
            if block.ordered:
                p = docx.add_paragraph()
                num_run = p.add_run(f"{i + 1}. ")
                num_run.font.color.rgb = RGBColor(*Colors.PRIMARY)
                num_run.font.size = Pt(Fonts.BODY_SIZE_PT)
            else:
                p = docx.add_paragraph()
                marker = p.add_run("●  ")
                marker.font.color.rgb = RGBColor(*Colors.PRIMARY)
                marker.font.size = Pt(8)

            # Use rich spans if available, otherwise plain text
            if block.rich_items and i < len(block.rich_items):
                _add_rich_runs(p, block.rich_items[i])
            else:
                text_run = p.add_run(item)
                text_run.font.size = Pt(Fonts.BODY_SIZE_PT)
                text_run.font.name = Fonts.BODY
            p.paragraph_format.space_after = Pt(3)
            p.paragraph_format.left_indent = Inches(0.4)

    def _render_blockquote(self, docx: DocxDocument, block: BlockquoteBlock) -> None:
        """Render a blockquote with left accent bar."""
        p = docx.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
        _set_paragraph_shading(p, Colors.BG_LIGHT)

        # Add left border via XML
        p_pr = p._element.get_or_add_pPr()
        p_bdr = OxmlElement("w:pBdr")
        left_bdr = OxmlElement("w:left")
        left_bdr.set(qn("w:val"), "single")
        left_bdr.set(qn("w:sz"), "24")
        left_bdr.set(qn("w:space"), "8")
        left_bdr.set(qn("w:color"), _hex(Colors.PRIMARY))
        p_bdr.append(left_bdr)
        p_pr.append(p_bdr)

        run = p.add_run(block.text)
        run.italic = True
        run.font.size = Pt(Fonts.BODY_SIZE_PT)
        run.font.color.rgb = RGBColor(*Colors.SECONDARY)
        run.font.name = Fonts.BODY

    # ------------------------------------------------------------------
    # Table renderer
    # ------------------------------------------------------------------

    def _render_table(self, docx: DocxDocument, block: TableBlock) -> None:
        cols = len(block.headers) if block.headers else (len(block.rows[0]) if block.rows else 0)
        if cols == 0:
            return

        # Add a small spacer before the table
        spacer = docx.add_paragraph()
        spacer.paragraph_format.space_after = Pt(2)
        spacer.paragraph_format.space_before = Pt(4)

        row_count = (1 if block.headers else 0) + len(block.rows)
        table = docx.add_table(rows=row_count, cols=cols)
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        table.style = "Table Grid"

        # Header row
        idx = 0
        if block.headers:
            for j, header in enumerate(block.headers):
                cell = table.cell(0, j)
                # Use rich spans if available (preserves hyperlinks)
                if block.rich_headers and j < len(block.rich_headers) and block.rich_headers[j]:
                    cell.text = ""  # clear default
                    p = cell.paragraphs[0]
                    _add_rich_runs(p, block.rich_headers[j])
                else:
                    cell.text = header
                _set_cell_shading(cell, Colors.TABLE_HEADER_BG)
                _set_cell_margins(cell, top=60, bottom=60, left=100, right=100)
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.bold = True
                        run.font.color.rgb = RGBColor(*Colors.WHITE)
                        run.font.size = Pt(10)
                        run.font.name = Fonts.HEADING
            idx = 1

        # Data rows with alternating colors
        for row_i, row_data in enumerate(block.rows):
            bg = Colors.TABLE_ALT_ROW if row_i % 2 == 0 else Colors.WHITE
            for j, val in enumerate(row_data):
                if j < cols:
                    cell = table.cell(idx, j)
                    # Use rich spans if available (preserves hyperlinks)
                    if (
                        block.rich_rows
                        and row_i < len(block.rich_rows)
                        and j < len(block.rich_rows[row_i])
                        and block.rich_rows[row_i][j]
                    ):
                        cell.text = ""  # clear default
                        p = cell.paragraphs[0]
                        _add_rich_runs(p, block.rich_rows[row_i][j])
                    else:
                        cell.text = val
                    _set_cell_shading(cell, bg)
                    _set_cell_margins(cell, top=40, bottom=40, left=100, right=100)
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.font.size = Pt(10)
                            run.font.name = Fonts.BODY
                            run.font.color.rgb = RGBColor(*Colors.TEXT)
            idx += 1

    # ------------------------------------------------------------------
    # Metadata page
    # ------------------------------------------------------------------

    def _add_metadata_page(self, docx: DocxDocument, doc: DocumentModel) -> None:
        docx.add_page_break()

        heading = docx.add_paragraph()
        run = heading.add_run("📎  Document Metadata")
        run.font.size = Pt(Fonts.H1_SIZE_PT)
        run.font.color.rgb = RGBColor(*Colors.PRIMARY_DARK)
        run.bold = True
        run.font.name = Fonts.HEADING
        _add_bottom_border(heading, Colors.ACCENT, size=8)

        docx.add_paragraph()

        meta = doc.metadata
        items = [
            ("Repository", meta.repo_name),
            ("URL", meta.repo_url),
            ("Source Path", meta.source_path),
            ("Generated At", meta.generated_at),
            ("Generator", "opendocs v0.1"),
            ("Sections", str(len(doc.sections))),
            ("Content Blocks", str(len(doc.all_blocks))),
            ("Mermaid Diagrams", str(len(doc.mermaid_diagrams))),
        ]

        meta_table = docx.add_table(rows=len(items), cols=2)
        meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, value) in enumerate(items):
            left = meta_table.cell(i, 0)
            right = meta_table.cell(i, 1)
            left.text = label
            right.text = value or "—"

            bg = Colors.BG_LIGHT if i % 2 == 0 else Colors.WHITE
            _set_cell_shading(left, bg)
            _set_cell_shading(right, bg)
            _set_cell_margins(left, top=50, bottom=50, left=100, right=100)
            _set_cell_margins(right, top=50, bottom=50, left=100, right=100)

            for p in left.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(*Colors.HEADING)
                    r.font.name = Fonts.BODY
            for p in right.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(*Colors.TEXT)
                    r.font.name = Fonts.BODY

    # ------------------------------------------------------------------
    # Knowledge Graph page
    # ------------------------------------------------------------------

    def _add_knowledge_graph_page(self, docx: DocxDocument) -> None:
        """Render a Knowledge Graph summary page in the Word document."""
        from ..core.knowledge_graph import EntityType

        kg = self.kg
        if not kg:
            return

        docx.add_page_break()

        heading = docx.add_paragraph()
        run = heading.add_run("Knowledge Graph Summary")
        run.font.size = Pt(Fonts.H1_SIZE_PT)
        run.font.color.rgb = RGBColor(*Colors.PRIMARY_DARK)
        run.bold = True
        run.font.name = Fonts.HEADING
        _add_bottom_border(heading, Colors.ACCENT, size=8)

        docx.add_paragraph()

        # Summary text
        if kg.summary:
            p = docx.add_paragraph()
            run = p.add_run(kg.summary)
            run.font.size = Pt(Fonts.BODY_SIZE_PT)
            run.font.name = Fonts.BODY
            docx.add_paragraph()

        # Executive summary (LLM-generated)
        if kg.executive_summary:
            exec_heading = docx.add_paragraph()
            run = exec_heading.add_run("Executive Summary")
            run.font.size = Pt(Fonts.H2_SIZE_PT)
            run.font.color.rgb = RGBColor(*Colors.PRIMARY_DARK)
            run.bold = True
            run.font.name = Fonts.HEADING

            p = docx.add_paragraph()
            _set_paragraph_shading(p, Colors.BG_LIGHT)
            run = p.add_run(kg.executive_summary)
            run.font.size = Pt(Fonts.BODY_SIZE_PT)
            run.font.name = Fonts.BODY
            run.font.color.rgb = RGBColor(*Colors.TEXT)
            docx.add_paragraph()

        # Stakeholder summaries (LLM-generated)
        if kg.stakeholder_summaries:
            stake_heading = docx.add_paragraph()
            run = stake_heading.add_run("👥  Stakeholder Views")
            run.font.size = Pt(Fonts.H2_SIZE_PT)
            run.font.color.rgb = RGBColor(*Colors.PRIMARY_DARK)
            run.bold = True
            run.font.name = Fonts.HEADING

            persona_labels = {
                "cto": "CTO / Technical Lead",
                "investor": "Investor / Business",
                "developer": "Developer Onboarding",
            }

            for persona, content in kg.stakeholder_summaries.items():
                if not content or content.startswith("["):
                    continue
                label = persona_labels.get(persona, persona.title())
                sub_h = docx.add_paragraph()
                run = sub_h.add_run(label)
                run.font.size = Pt(Fonts.H3_SIZE_PT)
                run.font.color.rgb = RGBColor(*Colors.SECONDARY)
                run.bold = True
                run.font.name = Fonts.HEADING

                for line in content.strip().split("\n"):
                    line = line.strip()
                    if not line:
                        continue
                    p = docx.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0.3)
                    if line.startswith(("- ", "• ", "* ")):
                        line = line[2:]
                        bullet = p.add_run("●  ")
                        bullet.font.color.rgb = RGBColor(*Colors.ACCENT)
                        bullet.font.size = Pt(8)
                    text_run = p.add_run(line)
                    text_run.font.size = Pt(Fonts.BODY_SIZE_PT)
                    text_run.font.name = Fonts.BODY

                docx.add_paragraph()

        # Stats table
        stats = kg.extraction_stats or kg.compute_stats()
        stat_items = [
            ("Total Entities", str(stats.get("total_entities", 0))),
            ("Total Relations", str(stats.get("total_relations", 0))),
            ("Deterministic", str(stats.get("deterministic_entities", 0))),
            ("LLM-Extracted", str(stats.get("llm_entities", 0))),
        ]
        stats_table = docx.add_table(rows=len(stat_items), cols=2)
        stats_table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for i, (label, val) in enumerate(stat_items):
            left = stats_table.cell(i, 0)
            right = stats_table.cell(i, 1)
            left.text = label
            right.text = val
            bg = Colors.BG_LIGHT if i % 2 == 0 else Colors.WHITE
            _set_cell_shading(left, bg)
            _set_cell_shading(right, bg)
            _set_cell_margins(left)
            _set_cell_margins(right)
            for p in left.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                for r in p.runs:
                    r.bold = True
                    r.font.size = Pt(10)
                    r.font.color.rgb = RGBColor(*Colors.HEADING)
            for p in right.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

        docx.add_paragraph()

        # Entity listing by type
        for et in EntityType:
            entities = kg.entities_of_type(et)
            if not entities:
                continue

            type_heading = docx.add_paragraph()
            run = type_heading.add_run(f"▸ {et.value.replace('_', ' ').title()}s ({len(entities)})")
            run.font.size = Pt(Fonts.H3_SIZE_PT)
            run.font.color.rgb = RGBColor(*Colors.HEADING)
            run.bold = True

            for entity in entities[:20]:  # Cap display
                p = docx.add_paragraph()
                p.paragraph_format.left_indent = Inches(0.4)
                marker = p.add_run("●  ")
                marker.font.color.rgb = RGBColor(*Colors.PRIMARY)
                marker.font.size = Pt(8)
                name_run = p.add_run(entity.name)
                name_run.bold = True
                name_run.font.size = Pt(Fonts.BODY_SIZE_PT)

                # Show key properties
                if entity.properties:
                    props_text = "  —  " + ", ".join(
                        f"{k}: {v}" for k, v in list(entity.properties.items())[:3]
                        if v and k not in ("url", "description")
                    )
                    if props_text.strip() != "—":
                        prop_run = p.add_run(props_text)
                        prop_run.font.size = Pt(Fonts.CAPTION_SIZE_PT)
                        prop_run.font.color.rgb = RGBColor(*Colors.MUTED)

                # Confidence indicator
                conf_run = p.add_run(f"  [{entity.confidence:.0%}]")
                conf_run.font.size = Pt(Fonts.SMALL_SIZE_PT)
                if entity.confidence >= 0.9:
                    conf_run.font.color.rgb = RGBColor(*Colors.SUCCESS)
                elif entity.confidence >= 0.7:
                    conf_run.font.color.rgb = RGBColor(*Colors.WARNING)
                else:
                    conf_run.font.color.rgb = RGBColor(*Colors.DANGER)

        # Mermaid graph from KG
        docx.add_paragraph()
        mermaid_heading = docx.add_paragraph()
        run = mermaid_heading.add_run("Auto-Generated Architecture Graph")
        run.font.size = Pt(Fonts.H2_SIZE_PT)
        run.font.color.rgb = RGBColor(*Colors.PRIMARY_DARK)
        run.bold = True

        # Try to embed rendered KG diagram image
        kg_img = self.image_cache.kg_diagram if self.image_cache else None
        if kg_img and kg_img.exists():
            docx.add_picture(str(kg_img), width=Inches(5.8))
            last_p = docx.paragraphs[-1]
            last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            last_p.paragraph_format.space_after = Pt(8)
        else:
            # Fallback: show mermaid code
            mermaid_code = kg.to_mermaid()
            code_p = docx.add_paragraph()
            _set_paragraph_shading(code_p, Colors.CODE_BG)
            for line in mermaid_code.split("\n"):
                run = code_p.add_run(line + "\n")
                run.font.name = Fonts.CODE
                run.font.size = Pt(Fonts.CODE_SIZE_PT)
                run.font.color.rgb = RGBColor(*Colors.CODE_TEXT)
